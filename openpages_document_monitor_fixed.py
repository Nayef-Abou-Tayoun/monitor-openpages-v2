#!/usr/bin/env python3
"""
OpenPages Document Monitor with Watson Orchestrate Integration
Monitors a specific process for new documents, downloads them to COS,
extracts text content, and triggers Watson Orchestrate executive summary agent

FIXED VERSION - Corrected Watson Orchestrate integration
"""

import requests
import time
import os
import json
import io
from datetime import datetime
import pytz
from requests.auth import HTTPBasicAuth
import ibm_boto3
from ibm_botocore.client import Config
from flask import Flask, jsonify
import threading

# Disable SSL warnings
requests.packages.urllib3.disable_warnings()

# Text extraction libraries
try:
    from PyPDF2 import PdfReader
    import docx
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠ Warning: PyPDF2 or python-docx not installed. Text extraction limited.")

# ============================================================================
# CONFIGURATION - Using environment variables
# ============================================================================

# OpenPages Server Configuration
OPENPAGES_SERVER = os.getenv("OPENPAGES_SERVER")
USERNAME = os.getenv("OPENPAGES_USERNAME")
PASSWORD = os.getenv("OPENPAGES_PASSWORD")

# Process to Monitor
PROCESS_ID = os.getenv("PROCESS_ID")
PROCESS_NAME = os.getenv("PROCESS_NAME")

# Monitoring Configuration
CHECK_INTERVAL_SECONDS = int(os.getenv("CHECK_INTERVAL_SECONDS", "60"))

# IBM Cloud Object Storage Configuration
COS_API_KEY = os.getenv("COS_API_KEY")
COS_INSTANCE_CRN = os.getenv("COS_INSTANCE_CRN")
COS_ENDPOINT = os.getenv("COS_ENDPOINT")
COS_BUCKET_NAME = os.getenv("COS_BUCKET_NAME")

# Watson Orchestrate Configuration
WXO_API_KEY = os.getenv("WXO_API_KEY")
WXO_INSTANCE_ID = os.getenv("WXO_INSTANCE_ID")
WXO_AGENT_ID = os.getenv("WXO_AGENT_ID", "99f19079-0e86-4baf-965f-a17ebc7e672b")

# ============================================================================
# DO NOT MODIFY BELOW THIS LINE
# ============================================================================

class OpenPagesDocumentMonitor:
    def __init__(self, server, username, password, process_id, check_interval,
                 cos_api_key=None, cos_instance_crn=None, cos_endpoint=None, cos_bucket=None,
                 wxo_api_key=None, wxo_instance_id=None, wxo_agent_id=None):
        # Normalize server URL - remove trailing slash for consistent path construction
        self.server = server.rstrip('/')
        self.username = username
        self.password = password
        self.process_id = process_id
        self.check_interval = check_interval
        self.known_documents = set()
        
        # Create a session for persistent connections and proper authentication
        self.session = requests.Session()
        self.session.auth = (username, password)  # Use tuple auth
        self.session.verify = False  # Disable SSL verification for OpenPages
        
        # Watson Orchestrate Configuration
        self.wxo_api_key = wxo_api_key
        self.wxo_instance_id = wxo_instance_id
        self.wxo_agent_id = wxo_agent_id
        self.wxo_token = None
        self.wxo_token_expiration = 0
        self.wxo_enabled = bool(wxo_api_key and wxo_instance_id and wxo_agent_id)
        
        # IBM COS Configuration
        self.cos_enabled = all([cos_api_key, cos_instance_crn, cos_endpoint, cos_bucket])
        self.cos_bucket = cos_bucket
        
        if self.cos_enabled:
            # Initialize IBM COS client
            self.cos_client = ibm_boto3.client(
                's3',
                ibm_api_key_id=cos_api_key,
                ibm_service_instance_id=cos_instance_crn,
                config=Config(signature_version='oauth'),
                endpoint_url=cos_endpoint
            )
            print(f"✓ IBM Cloud Object Storage client initialized")
            print(f"✓ Target bucket: {cos_bucket}")
            
            # Verify bucket exists and is accessible
            if self.verify_bucket_access():
                print(f"✓ Bucket verified and accessible")
            else:
                print(f"✗ ERROR: Cannot access bucket '{cos_bucket}'")
                self.cos_enabled = False
        else:
            self.cos_client = None
            print(f"⚠ WARNING: IBM Cloud Object Storage not configured!")
        
        if self.wxo_enabled:
            print(f"✓ Watson Orchestrate agent integration enabled")
            print(f"✓ Instance ID: {wxo_instance_id}")
            print(f"✓ Agent ID: {wxo_agent_id}")
        else:
            print(f"⚠ Watson Orchestrate agent not configured")
    
    def verify_bucket_access(self):
        """Verify that the bucket exists and is accessible"""
        try:
            test_key = "test/connection_test.txt"
            test_content = b"Connection test"
            self.cos_client.put_object(
                Bucket=self.cos_bucket,
                Key=test_key,
                Body=test_content
            )
            try:
                self.cos_client.delete_object(Bucket=self.cos_bucket, Key=test_key)
            except:
                pass
            return True
        except Exception as e:
            print(f"   Error details: {str(e)}")
            return False
    
    def log(self, message):
        """Print timestamped log message"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print(f"[{timestamp}] {message}")
    
    def extract_text_from_file(self, file_content, filename):
        """Extract text content from various file types"""
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            # Text files
            if ext in ['.txt', '.csv', '.json', '.xml', '.log']:
                return file_content.decode('utf-8', errors='ignore')
            
            # PDF files
            elif ext == '.pdf' and PDF_SUPPORT:
                pdf_file = io.BytesIO(file_content)
                reader = PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            
            # Word documents
            elif ext in ['.docx'] and PDF_SUPPORT:
                doc_file = io.BytesIO(file_content)
                doc = docx.Document(doc_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text
            
            else:
                return f"[Binary file: {filename} - text extraction not supported for {ext}]"
        
        except Exception as e:
            self.log(f"⚠ Error extracting text from {filename}: {str(e)}")
            return f"[Error extracting text from {filename}]"
    
    def is_token_expired(self, expiration_time):
        """Check if the token is expired (with 60 second buffer)"""
        return int(time.time()) >= (expiration_time - 60)
    
    def get_bearer_token(self):
        """Get or refresh the Watson Orchestrate bearer token - FIXED VERSION"""
        if self.wxo_token and not self.is_token_expired(self.wxo_token_expiration):
            return self.wxo_token
        
        try:
            # FIXED: Correct IBM Cloud IAM token endpoint
            url = "https://iam.cloud.ibm.com/identity/token"
            
            headers = {
                "Content-Type": "application/x-www-form-urlencoded",
                "Accept": "application/json"
            }
            
            data = {
                "grant_type": "urn:ibm:params:oauth:grant-type:apikey",
                "apikey": self.wxo_api_key
            }
            
            self.log(f"🔑 Requesting Watson Orchestrate IAM token...")
            response = requests.post(url, headers=headers, data=data, timeout=30)
            
            if response.status_code == 200:
                token_data = response.json()
                self.wxo_token = token_data.get("access_token")
                expires_in = token_data.get("expires_in", 3600)
                self.wxo_token_expiration = int(time.time()) + expires_in
                self.log(f"✓ Watson Orchestrate token obtained (expires in {expires_in}s)")
                return self.wxo_token
            else:
                self.log(f"⚠ Token request failed: HTTP {response.status_code}")
                self.log(f"   Response: {response.text[:500]}")
                return None
        except Exception as e:
            self.log(f"⚠ Exception getting token: {str(e)}")
            import traceback
            self.log(f"   Traceback: {traceback.format_exc()}")
            return None
    
    def trigger_wxo_executive_summary(self, document_text, filename, doc_id):
        """Trigger Watson Orchestrate executive summary agent - FIXED VERSION"""
        if not self.wxo_enabled:
            return None
        
        try:
            # Get bearer token
            token = self.get_bearer_token()
            if not token:
                self.log(f"⚠ Failed to get Watson Orchestrate token")
                return None
            
            # Prepare the prompt for the agent
            prompt = f"""Please analyze this document and provide an executive summary:

Document Name: {filename}
Document ID: {doc_id}
Process ID: {self.process_id}
Process Name: {PROCESS_NAME}

Document Content:
{document_text[:8000]}

Please provide:
1. Executive Summary
2. Key Findings
3. Risk Assessment
4. Recommendations
"""
            
            # FIXED: Correct Watson Orchestrate API endpoint format
            url = f"https://api.watsonx.ai/v1/watson_orchestrate/instances/{self.wxo_instance_id}/agents/{self.wxo_agent_id}/chat/completions"
            
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
                "Accept": "application/json"
            }
            
            payload = {
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "stream": False
            }
            
            self.log(f"🤖 Triggering Watson Orchestrate executive summary agent...")
            self.log(f"   Endpoint: {url}")
            
            response = requests.post(url, headers=headers, json=payload, timeout=120)
            
            if response.status_code == 200:
                data = response.json()
                summary = data.get("choices", [{}])[0].get("message", {}).get("content", "No response returned.")
                self.log(f"✅ Executive summary generated successfully!")
                self.log(f"   Summary preview: {summary[:200]}...")
                
                # Use EST timezone for timestamp
                est = pytz.timezone('US/Eastern')
                est_time = datetime.now(est)
                
                result = {
                    "summary": summary,
                    "document_name": filename,
                    "document_id": doc_id,
                    "process_id": self.process_id,
                    "timestamp": est_time.strftime("%d/%m/%Y %H:%M:%S")
                }
                return result
            else:
                self.log(f"⚠ Watson Orchestrate agent failed: HTTP {response.status_code}")
                self.log(f"   Response: {response.text[:500]}")
                
                # If unauthorized, clear token to force refresh on next attempt
                if response.status_code == 401:
                    self.log(f"   Token expired or invalid - will refresh on next attempt")
                    self.wxo_token = None
                    self.wxo_token_expiration = 0
                
                return None
        
        except Exception as e:
            self.log(f"⚠ Exception calling Watson Orchestrate agent: {str(e)}")
            import traceback
            self.log(f"   Traceback: {traceback.format_exc()}")
            return None
    
    def get_process_documents(self):
        """Get all documents associated with the process"""
        try:
            url = f"{self.server}/grc/api/contents/{self.process_id}/associations/children"
            headers = {"Accept": "application/json"}
            
            response = self.session.get(url, headers=headers, timeout=30)
            
            if response.status_code == 200:
                # Check if response has content
                if not response.text or response.text.strip() == '':
                    self.log(f"⚠ API returned empty response (HTTP 200 but no content)")
                    return []
                
                children = response.json()
                documents = [child for child in children if child.get('typeDefinitionId') in ['22', '42', '46']]
                return documents
            else:
                self.log(f"⚠ Error getting documents: HTTP {response.status_code}")
                self.log(f"   Response: {response.text[:200]}")
                return []
        except requests.exceptions.JSONDecodeError as e:
            self.log(f"⚠ JSON decode error: {str(e)}")
            self.log(f"   Response text: {response.text[:200] if response else 'No response'}")
            return []
        except requests.exceptions.RequestException as e:
            self.log(f"⚠ Request error: {str(e)}")
            return []
        except Exception as e:
            self.log(f"⚠ Exception getting documents: {str(e)}")
            return []
    
    def get_document_details(self, doc_id):
        """Get detailed information about a document"""
        try:
            url = f"{self.server}/grc/api/contents/{doc_id}"
            
            response = self.session.get(url, timeout=30)
            
            if response.status_code == 200:
                return response.json()
            else:
                return None
        except Exception as e:
            self.log(f"⚠ Exception getting document details: {str(e)}")
            return None
    
    def upload_to_cos_direct(self, file_content, filename):
        """Upload file content directly to IBM Cloud Object Storage"""
        if not self.cos_enabled:
            return False
        
        try:
            object_key = f"Process_{self.process_id}/{filename}"
            
            self.cos_client.put_object(
                Bucket=self.cos_bucket,
                Key=object_key,
                Body=file_content
            )
            
            self.log(f"☁️  Uploaded to COS: {object_key} ({len(file_content):,} bytes)")
            return True
        except Exception as e:
            self.log(f"⚠ Exception uploading to COS: {str(e)}")
            return False
    
    def format_summary_as_docx(self, wxo_result, filename, doc_id):
        """Format Watson Orchestrate summary as a DOCX document"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re
        except ImportError:
            self.log("⚠ python-docx not available, skipping DOCX generation")
            return None
        
        doc = Document()
        
        # Title
        title = doc.add_heading('WATSON ORCHESTRATE EXECUTIVE SUMMARY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Document Information Section
        doc.add_heading('Document Information', level=1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Document Name:', filename),
            ('Document ID:', str(doc_id)),
            ('Process ID:', str(self.process_id)),
            ('Process Name:', PROCESS_NAME),
            ('Generated:', wxo_result.get('timestamp', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            # Bold the labels
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Add spacing
        
        # Executive Summary Section
        doc.add_heading('Executive Summary', level=1)
        
        summary = wxo_result.get('summary', 'No summary available')
        
        # Parse and format the summary
        lines = summary.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                current_paragraph = None
                continue
            
            # Main headers (starts with ** and ends with ** and no other text)
            if line.startswith('**') and line.endswith('**') and line.count('**') == 2:
                header_text = line.replace('**', '').strip()
                doc.add_heading(header_text, level=2)
                current_paragraph = None
            # Numbered items
            elif len(line) > 2 and line[0:2].replace('.', '').replace(')', '').isdigit():
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, line)
                current_paragraph = None
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, line[2:])
                current_paragraph = None
            # Regular text
            else:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                    self._add_formatted_text(current_paragraph, line)
                else:
                    current_paragraph.add_run(' ')
                    self._add_formatted_text(current_paragraph, line)
        
        # Save to bytes
        import io
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    
    def _add_formatted_text(self, paragraph, text):
        """Add text to paragraph with inline bold formatting support"""
        import re
        
        # Split text by ** markers for bold formatting
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                bold_text = part[2:-2]  # Remove ** markers
                run = paragraph.add_run(bold_text)
                run.font.bold = True
            elif part:
                # Regular text
                paragraph.add_run(part)
    
    def get_document_folder_id(self):
        """
        Get the folder ID where documents associated with this process are stored.
        This checks existing child documents to find the correct folder.
        """
        try:
            # Get existing child documents of the process
            url = f"{self.server}/grc/api/contents/{self.process_id}/associations/children"
            
            response = self.session.get(url, timeout=30)
            
            if response.status_code == 200:
                children = response.json()
                
                # Find a document (type 22) to get its folder
                for child in children:
                    if child.get('typeDefinitionId') == '22':  # Document type
                        doc_id = child.get('id')
                        
                        # Get this document's details to find its folder
                        doc_url = f"{self.server}/grc/api/contents/{doc_id}"
                        doc_response = self.session.get(doc_url, timeout=30)
                        
                        if doc_response.status_code == 200:
                            doc_data = doc_response.json()
                            folder_id = doc_data.get('parentFolderId')
                            doc_path = doc_data.get('path', '')
                            
                            if folder_id:
                                self.log(f"✓ Found document folder ID: {folder_id}")
                                self.log(f"  (from existing document: {doc_path})")
                                return folder_id
                
                self.log(f"⚠ No existing documents found to determine folder")
            else:
                self.log(f"⚠ Failed to get process children: HTTP {response.status_code}")
            
            return None
            
        except Exception as e:
            self.log(f"⚠ Exception getting document folder: {str(e)}")
            return None
    
    def upload_document_to_openpages(self, file_content, filename, description="Executive Risk Summary"):
        """Upload a document back to OpenPages and associate it with the process"""
        try:
            # Get the folder ID where documents for this process are stored
            folder_id = self.get_document_folder_id()
            
            if not folder_id:
                self.log(f"⚠ Cannot upload: Unable to determine document folder for process {self.process_id}")
                return None
            
            # Encode file content to base64 for binary files
            import base64
            file_content_b64 = base64.b64encode(file_content).decode('utf-8')
            
            # Step 1: Create document in the same folder as other process documents
            create_url = f"{self.server}/grc/api/contents"
            
            create_payload = {
                "contentDefinition": {
                    "attribute": {
                        "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    },
                    "children": file_content_b64
                },
                "fileTypeDefinition": {
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "fileExtension": "docx"
                },
                "fields": {
                    "field": []
                },
                "typeDefinitionId": "22",  # Document type
                "parentFolderId": folder_id,  # Use the documents folder
                "name": filename,
                "description": description
            }
            
            self.log(f"📤 Step 1: Creating document in folder {folder_id}: {filename}")
            create_response = self.session.post(create_url, json=create_payload, timeout=60)
            
            if create_response.status_code not in [200, 201]:
                self.log(f"⚠ Failed to create document: HTTP {create_response.status_code}")
                self.log(f"   Response: {create_response.text[:500]}")
                return None
            
            doc_data = create_response.json()
            new_doc_id = doc_data.get('id')
            self.log(f"✓ Document created with ID: {new_doc_id}")
            
            # Step 2: Associate document as child of the process
            assoc_url = f"{self.server}/grc/api/contents/{self.process_id}/associations/children"
            assoc_payload = [{"id": new_doc_id}]
            
            self.log(f"🔗 Step 2: Associating document with process {self.process_id}...")
            assoc_response = self.session.post(assoc_url, json=assoc_payload, timeout=30)
            
            if assoc_response.status_code not in [200, 201, 204]:
                self.log(f"⚠ Failed to associate: HTTP {assoc_response.status_code}")
                self.log(f"   Response: {assoc_response.text[:500]}")
                self.log(f"   Document {new_doc_id} created but not visible in process Admin tab")
                return new_doc_id
            
            self.log(f"✓ Document associated with process successfully")
            self.log(f"🎉 Executive summary uploaded to OpenPages: {filename} (ID: {new_doc_id})")
            self.log(f"   Document will appear in Process {self.process_id} Admin tab")
            
            return new_doc_id
            
        except Exception as e:
            self.log(f"⚠ Exception uploading document to OpenPages: {str(e)}")
            import traceback
            self.log(f"   Traceback: {traceback.format_exc()}")
            return None
    
    def download_and_process_document(self, doc_id, filename):
        """Download document, upload to COS, extract text, and trigger WXO agent"""
        try:
            url = f"{self.server}/grc/api/contents/{doc_id}/document"
            
            self.session.headers.update({"Accept": "application/octet-stream"})
            response = self.session.get(url, timeout=60, stream=True)
            self.session.headers.update({"Accept": "application/json"})
            
            if response.status_code == 200:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                name, ext = os.path.splitext(filename)
                unique_filename = f"{name}_{timestamp}{ext}"
                
                # Read content into memory
                file_content = b''
                for chunk in response.iter_content(chunk_size=8192):
                    file_content += chunk
                
                file_size = len(file_content)
                self.log(f"✓ Downloaded: {unique_filename} ({file_size:,} bytes)")
                
                # Upload to IBM Cloud Object Storage
                if self.cos_enabled:
                    self.upload_to_cos_direct(file_content, unique_filename)
                
                # Extract text content
                self.log(f"📄 Extracting text content from {unique_filename}...")
                document_text = self.extract_text_from_file(file_content, filename)
                text_length = len(document_text)
                self.log(f"✓ Extracted {text_length:,} characters of text")
                
                # Trigger Watson Orchestrate executive summary agent
                if self.wxo_enabled and text_length > 50:
                    wxo_result = self.trigger_wxo_executive_summary(document_text, filename, doc_id)
                    if wxo_result:
                        # Save the summary as JSON to COS
                        summary_filename_json = f"{name}_summary_{timestamp}.json"
                        summary_content_json = json.dumps(wxo_result, indent=2).encode('utf-8')
                        if self.cos_enabled:
                            self.upload_to_cos_direct(summary_content_json, summary_filename_json)
                        
                        # Save the summary as formatted DOCX document to COS
                        summary_filename_docx = f"{name}_summary_{timestamp}.docx"
                        formatted_doc_bytes = self.format_summary_as_docx(wxo_result, filename, doc_id)
                        if formatted_doc_bytes:
                            if self.cos_enabled:
                                self.upload_to_cos_direct(formatted_doc_bytes, summary_filename_docx)
                                self.log(f"📄 Formatted DOCX summary saved to COS: {summary_filename_docx}")
                            
                            # Upload the executive summary back to OpenPages
                            base_name = name.replace('.docx', '').replace('.txt', '').replace('.pdf', '')
                            openpages_doc_name = f"Executive Risk Summary - {base_name}"
                            uploaded_doc_id = self.upload_document_to_openpages(
                                formatted_doc_bytes,
                                openpages_doc_name,
                                f"AI-generated executive summary for {filename}"
                            )
                            if uploaded_doc_id:
                                self.log(f"✅ Executive summary uploaded to OpenPages (Doc ID: {uploaded_doc_id})")
                
                return True
            else:
                self.log(f"⚠ Failed to download document {doc_id}: HTTP {response.status_code}")
                return False
        except Exception as e:
            self.log(f"⚠ Exception processing document: {str(e)}")
            import traceback
            self.log(f"   Traceback: {traceback.format_exc()}")
            return False
    
    def check_for_new_documents(self):
        """Check for new documents and process them"""
        documents = self.get_process_documents()
        
        if not documents:
            self.log(f"ℹ No documents found in process {self.process_id}")
            return
        
        new_docs_found = False
        
        for doc in documents:
            doc_id = doc.get('id')
            
            if doc_id not in self.known_documents:
                new_docs_found = True
                self.log(f"🆕 NEW DOCUMENT DETECTED: ID {doc_id}")
                
                details = self.get_document_details(doc_id)
                
                if details:
                    path = details.get('path', '')
                    filename = os.path.basename(path) if path else details.get('name', f'document_{doc_id}.bin')
                    
                    self.log(f"   Filename: {filename}")
                    self.log(f"   Path: {path}")
                    
                    # Skip processing our own generated executive summaries
                    if filename.startswith("Executive Risk Summary"):
                        self.log(f"   ⏭ Skipping - this is a generated summary")
                        self.known_documents.add(doc_id)
                        continue
                    
                    if self.download_and_process_document(doc_id, filename):
                        self.known_documents.add(doc_id)
                        self.log(f"✓ Successfully processed document {doc_id}")
                    else:
                        self.log(f"✗ Failed to process document {doc_id}")
                else:
                    self.log(f"⚠ Could not get details for document {doc_id}")
        
        if not new_docs_found:
            self.log(f"ℹ No new documents (currently tracking {len(self.known_documents)} documents)")
    
    def initialize_known_documents(self):
        """Download all existing documents, then start monitoring for new ones"""
        self.log("Initializing - scanning for existing documents...")
        documents = self.get_process_documents()
        
        if not documents:
            self.log("ℹ No existing documents found in process")
            self.log("✓ Monitoring started - will check every 60 seconds")
            return
        
        self.log(f"📥 Found {len(documents)} existing document(s) - downloading all...")
        self.log("-" * 70)
        
        for i, doc in enumerate(documents, 1):
            doc_id = doc.get('id')
            self.log(f"Processing existing document {i}/{len(documents)}: ID {doc_id}")
            
            details = self.get_document_details(doc_id)
            if details:
                path = details.get('path', '')
                filename = os.path.basename(path) if path else details.get('name', f'document_{doc_id}.bin')
                
                # Skip processing our own generated executive summaries
                if filename.startswith("Executive Risk Summary"):
                    self.log(f"   ⏭ Skipping - this is a generated summary")
                    self.known_documents.add(doc_id)
                    continue
                
                if self.download_and_process_document(doc_id, filename):
                    self.known_documents.add(doc_id)
                    self.log(f"✓ Successfully processed existing document {doc_id}")
                else:
                    self.log(f"✗ Failed to process existing document {doc_id}")
            
            self.log("-" * 70)
        
        self.log(f"✓ Downloaded {len(self.known_documents)} existing document(s)")
        self.log(f"✓ Now monitoring for NEW documents only - will check every {self.check_interval} seconds")
    
    def start_monitoring(self):
        """Start the monitoring loop"""
        self.log("=" * 70)
        self.log("OpenPages Document Monitor Started")
        self.log("=" * 70)
        self.log(f"Server: {self.server}")
        self.log(f"Process ID: {self.process_id}")
        self.log(f"Check Interval: {self.check_interval} seconds")
        self.log(f"COS Bucket: {self.cos_bucket}")
        self.log(f"COS Upload: {'ENABLED ☁️' if self.cos_enabled else 'DISABLED'}")
        self.log(f"Watson Orchestrate: {'ENABLED 🤖' if self.wxo_enabled else 'DISABLED'}")
        self.log(f"Storage: Direct to Cloud (no local files)")
        self.log("=" * 70)
        
        # Initialize by downloading existing documents
        self.initialize_known_documents()
        
        self.log(" ")
        self.log("👀 Now monitoring for new documents... (Press Ctrl+C to stop)")
        self.log(" ")
        
        # Start monitoring loop
        while True:
            try:
                self.check_for_new_documents()
                time.sleep(self.check_interval)
            except KeyboardInterrupt:
                self.log("\n⏹ Monitoring stopped by user")
                break
            except Exception as e:
                self.log(f"⚠ Error in monitoring loop: {str(e)}")
                time.sleep(self.check_interval)

# Flask app for health checks
app = Flask(__name__)
monitor = None

@app.route('/health')
def health():
    return jsonify({"status": "healthy", "service": "openpages-document-monitor"}), 200

@app.route('/status')
def status():
    if monitor:
        return jsonify({
            "status": "running",
            "process_id": monitor.process_id,
            "known_documents": len(monitor.known_documents),
            "cos_enabled": monitor.cos_enabled,
            "wxo_enabled": monitor.wxo_enabled
        }), 200
    return jsonify({"status": "not_initialized"}), 503

def run_flask():
    """Run Flask app in a separate thread"""
    app.run(host='0.0.0.0', port=8080, debug=False, use_reloader=False)

if __name__ == "__main__":
    # Start Flask health check server in background
    flask_thread = threading.Thread(target=run_flask, daemon=True)
    flask_thread.start()
    
    print("🌐 Starting HTTP server on port 8080...")
    print("   Health check: http://localhost:8080/health")
    print("   Status: http://localhost:8080/status")
    print(" ")
    
    # Create and start monitor
    monitor = OpenPagesDocumentMonitor(
        server=OPENPAGES_SERVER,
        username=USERNAME,
        password=PASSWORD,
        process_id=PROCESS_ID,
        check_interval=CHECK_INTERVAL_SECONDS,
        cos_api_key=COS_API_KEY,
        cos_instance_crn=COS_INSTANCE_CRN,
        cos_endpoint=COS_ENDPOINT,
        cos_bucket=COS_BUCKET_NAME,
        wxo_api_key=WXO_API_KEY,
        wxo_instance_id=WXO_INSTANCE_ID,
        wxo_agent_id=WXO_AGENT_ID
    )
    
    monitor.start_monitoring()

# Fixed by Bob - Watson Orchestrate Integration Corrected

# Made with Bob
