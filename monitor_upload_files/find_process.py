#!/usr/bin/env python3
"""
OpenPages Process Finder with Watson Orchestrate Integration
Finds processes by ID pattern and processes documents with AI analysis
"""

import os
import sys
import asyncio
import argparse
import base64
import httpx
import requests
import json
import io
import time
import re
from datetime import datetime
from typing import Optional, List, Dict, Any
from dotenv import load_dotenv
from aiohttp import web
import threading
import pytz

# Text extraction libraries
try:
    from PyPDF2 import PdfReader
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠ Warning: PyPDF2 or python-docx not installed. Text extraction limited.")

# IBM COS
try:
    import ibm_boto3
    from ibm_botocore.client import Config
    COS_SUPPORT = True
except ImportError:
    COS_SUPPORT = False
    print("⚠ Warning: ibm-cos-sdk not installed. COS upload disabled.")

# Disable SSL warnings
requests.packages.urllib3.disable_warnings()

# Load .env file from the same directory as this script
script_dir = os.path.dirname(os.path.abspath(__file__))
env_path = os.path.join(script_dir, '.env')
if os.path.exists(env_path):
    load_dotenv(env_path)
    print(f"✓ Loaded configuration from: {env_path}\n")
else:
    print(f"⚠ Warning: .env file not found at {env_path}")
    print(f"   Using environment variables or defaults\n")


# Simplified OpenPages Client (embedded)
class OpenPagesClient:
    """Simplified OpenPages API client for basic authentication"""
    
    def __init__(self, base_url: str, username: str, password: str, **kwargs):
        self.base_url = base_url.rstrip('/')
        self.username = username
        self.password = password
        self._http_client = None
        
    async def initialize_auth(self):
        """Initialize authentication (no-op for basic auth)"""
        pass
    
    async def close(self):
        """Close HTTP client"""
        if self._http_client:
            await self._http_client.aclose()
    
    async def query(self, query: str, limit: int = 100) -> Dict[str, Any]:
        """Execute a query against OpenPages API"""
        if not self._http_client:
            self._http_client = httpx.AsyncClient(verify=False, follow_redirects=True, timeout=30.0)
        
        url = f"{self.base_url}/opgrc/api/v2/query"
        auth_header = base64.b64encode(f"{self.username}:{self.password}".encode()).decode()
        
        headers = {
            'Authorization': f'Basic {auth_header}',
            'Content-Type': 'application/json',
            'Accept': 'application/json'
        }
        
        payload = {
            'statement': query,
            'offset': 0,
            'max_rows': 500,
            'limit': limit,
            'case_insensitive': False,
            'honor_primary': False
        }
        
        try:
            response = await self._http_client.post(url, json=payload, headers=headers)
            
            if response.status_code == 200:
                return response.json()
            else:
                print(f"⚠ Query failed with status {response.status_code}")
                return {'rows': []}
        except Exception as e:
            print(f"⚠ Request error: {str(e)}")
            return {'rows': []}


# Simple settings class
class Settings:
    def __init__(self):
        self.OPENPAGES_BASE_URL = os.getenv('OPENPAGES_BASE_URL', '')
        self.OPENPAGES_AUTHENTICATION_TYPE = os.getenv('OPENPAGES_AUTHENTICATION_TYPE', 'basic')
        self.OPENPAGES_USERNAME = os.getenv('OPENPAGES_USERNAME', '')
        self.OPENPAGES_PASSWORD = os.getenv('OPENPAGES_PASSWORD', '')
        self.OPENPAGES_APIKEY = os.getenv('OPENPAGES_APIKEY', '')
        self.OPENPAGES_AUTHENTICATION_URL = os.getenv('OPENPAGES_AUTHENTICATION_URL', '')

settings = Settings()


class ProcessFinder:
    """Find OpenPages processes by ID pattern"""
    
    def __init__(self, client: OpenPagesClient):
        self.client = client
        self.cos_client = None
        self.cos_bucket = None
        
        # Watson Orchestrate Configuration
        self.wxo_api_key = os.getenv('WXO_API_KEY')
        self.wxo_instance_id = os.getenv('WXO_INSTANCE_ID')
        self.wxo_agent_id = os.getenv('WXO_AGENT_ID', '99f19079-0e86-4baf-965f-a17ebc7e672b')
        self.wxo_token = None
        self.wxo_token_expiration = 0
        self.wxo_enabled = bool(self.wxo_api_key and self.wxo_instance_id and self.wxo_agent_id)
        
        # Initialize COS if credentials are available
        cos_api_key = os.getenv('COS_API_KEY')
        cos_instance_crn = os.getenv('COS_INSTANCE_CRN')
        cos_endpoint = os.getenv('COS_ENDPOINT')
        cos_bucket = os.getenv('COS_BUCKET_NAME')
        
        if all([cos_api_key, cos_instance_crn, cos_endpoint, cos_bucket]) and COS_SUPPORT:
            try:
                import ibm_boto3
                from ibm_botocore.client import Config
                
                self.cos_client = ibm_boto3.client(
                    's3',
                    ibm_api_key_id=cos_api_key,
                    ibm_service_instance_id=cos_instance_crn,
                    config=Config(signature_version='oauth'),
                    endpoint_url=cos_endpoint
                )
                self.cos_bucket = cos_bucket
                print(f"✓ IBM Cloud Object Storage initialized")
                print(f"   Bucket: {cos_bucket}")
            except Exception as e:
                print(f"⚠️  Failed to initialize COS: {str(e)}")
        
        if self.wxo_enabled:
            print(f"✓ Watson Orchestrate integration enabled")
            print(f"   Instance ID: {self.wxo_instance_id}")
            print(f"   Agent ID: {self.wxo_agent_id}")
        else:
            print(f"⚠ Watson Orchestrate not configured")
        
        print()  # Empty line for spacing
    
    async def find_process_by_id(self, process_id: str) -> Optional[Dict[str, Any]]:
        """
        Find a specific process by its exact ID
        
        Args:
            process_id: The process ID (e.g., "AML_PROC_0008")
            
        Returns:
            Process details if found, None otherwise
        """
        try:
            print(f"🔍 Searching for process: {process_id}")
            
            # Query for the specific process by Name (which contains the process ID)
            query = f"SELECT * FROM [SOXProcess] WHERE [Name] = '{process_id}'"
            
            result = await self.client.query(query, limit=1)
            
            if result and result.get('rows'):
                process = result['rows'][0]
                print(f"✅ Found process: {process_id}")
                return process
            else:
                print(f"❌ Process not found: {process_id}")
                return None
                
        except Exception as e:
            print(f"❌ Error searching for process: {str(e)}")
            return None
    
    async def find_processes_by_pattern(self, pattern: str) -> List[Dict[str, Any]]:
        """
        Find processes matching a pattern (e.g., "AML_PROC_*")
        
        Args:
            pattern: Pattern to match (supports wildcards)
            
        Returns:
            List of matching processes
        """
        try:
            print(f"🔍 Searching for processes matching pattern: {pattern}")
            
            # Convert wildcard pattern to SQL LIKE pattern
            sql_pattern = pattern.replace('*', '%').replace('?', '_')
            
            # Query for processes matching the pattern using Name field
            query = f"SELECT * FROM [SOXProcess] WHERE [Name] LIKE '{sql_pattern}'"
            
            result = await self.client.query(query, limit=100)
            
            if result and result.get('rows'):
                processes = result['rows']
                print(f"✅ Found {len(processes)} process(es) matching pattern")
                return processes
            else:
                print(f"❌ No processes found matching pattern: {pattern}")
                return []
                
        except Exception as e:
            print(f"❌ Error searching for processes: {str(e)}")
            return []
    
    async def find_all_processes(self, limit: int = 50) -> List[Dict[str, Any]]:
        """
        Find all processes (up to limit)
        
        Args:
            limit: Maximum number of processes to return
            
        Returns:
            List of processes
        """
        try:
            print(f"🔍 Retrieving all processes (limit: {limit})")
            
            query = "SELECT * FROM [SOXProcess]"
            
            result = await self.client.query(query, limit=limit)
            
            if result and result.get('rows'):
                processes = result['rows']
                print(f"✅ Found {len(processes)} process(es)")
                return processes
            else:
                print(f"❌ No processes found")
                return []
                
        except Exception as e:
            print(f"❌ Error retrieving processes: {str(e)}")
            return []
    
    async def find_process_documents(self, process_id: str) -> List[Dict[str, Any]]:
        """
        Find all documents (uploaded files) associated with a process
        
        Args:
            process_id: The process ID (e.g., "AML_PROC_0008")
            
        Returns:
            List of documents
        """
        try:
            print(f"📄 Searching for documents in process: {process_id}")
            
            # First, get the process to find its Resource ID
            process_query = f"SELECT * FROM [SOXProcess] WHERE [Name] = '{process_id}'"
            process_result = await self.client.query(process_query, limit=1)
            
            if not process_result or not process_result.get('rows'):
                print(f"❌ Process not found: {process_id}")
                return []
            
            process = process_result['rows'][0]
            resource_id = self._get_field_value(process, 'Resource ID')
            
            print(f"   Process Resource ID: {resource_id}")
            
            # Store resource_id as instance variable for later use in upload
            self.process_resource_id = resource_id
            
            # Use the OpenPages API to get child associations (documents are children of processes)
            # This is more reliable than complex queries
            try:
                import httpx
                # Remove /openpages from base_url if present and add /grc/api
                base = self.client.base_url.replace('/openpages', '').rstrip('/')
                url = f"{base}/grc/api/contents/{resource_id}/associations/children"
                headers = {"Accept": "application/json"}
                
                # Create a simple HTTP client for this request with redirect following
                async with httpx.AsyncClient(verify=False, follow_redirects=True) as http_client:
                    response = await http_client.get(
                        url,
                        headers=headers,
                        auth=(self.client.username, self.client.password),
                        timeout=30.0
                    )
                    
                    if response.status_code == 200:
                        # Check if response has content
                        if not response.text or response.text.strip() == '':
                            print(f"ℹ️  No documents found for process: {process_id} (empty response)")
                            return []
                        
                        try:
                            children = response.json()
                        except Exception as json_error:
                            print(f"⚠️  Response is not JSON. Status: {response.status_code}")
                            print(f"   Response text: {response.text[:200]}")
                            return []
                        
                        # Filter for documents (type 4 for this system, also check 22, 42, 46)
                        documents = [child for child in children if child.get('typeDefinitionId') in ['4', '22', '42', '46']]
                        
                        if documents:
                            print(f"✅ Found {len(documents)} document(s)")
                            print(f"   Fetching detailed information for each document...")
                            
                            # Fetch full details for each document
                            formatted_docs = []
                            for i, doc in enumerate(documents, 1):
                                doc_id = doc.get('id')
                                doc_path = doc.get('path', 'Unknown')
                                
                                # Extract filename from path
                                filename = doc_path.split('/')[-1] if doc_path != 'Unknown' else 'Unknown'
                                
                                # Fetch full document details
                                try:
                                    detail_url = f"{base}/grc/api/contents/{doc_id}"
                                    detail_response = await http_client.get(
                                        detail_url,
                                        headers=headers,
                                        auth=(self.client.username, self.client.password),
                                        timeout=30.0
                                    )
                                    
                                    if detail_response.status_code == 200:
                                        doc_details = detail_response.json()
                                        doc_name = doc_details.get('name', filename)
                                        doc_desc = doc_details.get('description', '')
                                        created_date = doc_details.get('createdDate', 'N/A')
                                    else:
                                        doc_name = filename
                                        doc_desc = ''
                                        created_date = 'N/A'
                                except Exception as e:
                                    doc_name = filename
                                    doc_desc = ''
                                    created_date = 'N/A'
                                
                                formatted_doc = {
                                    'fields': [
                                        {'name': 'Resource ID', 'value': doc_id},
                                        {'name': 'Name', 'value': doc_name},
                                        {'name': 'Description', 'value': doc_desc},
                                        {'name': 'Location', 'value': doc_path},
                                        {'name': 'Creation Date', 'value': created_date},
                                    ]
                                }
                                formatted_docs.append(formatted_doc)
                                print(f"      [{i}/{len(documents)}] {doc_name}")
                            
                            return formatted_docs
                        else:
                            print(f"ℹ️  No documents found for process: {process_id}")
                            return []
                    else:
                        print(f"⚠️  API returned status {response.status_code}")
                        return []
                        
            except Exception as api_error:
                print(f"❌ Error calling OpenPages API: {str(api_error)}")
                return []
                
        except Exception as e:
            print(f"❌ Error searching for documents: {str(e)}")
            import traceback
            traceback.print_exc()
            return []
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text content from various file types"""
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            # Text files
            if ext in ['.txt', '.csv', '.json', '.xml', '.log']:
                return file_content.decode('utf-8', errors='ignore')
            
            # PDF files
            elif ext == '.pdf' and PDF_SUPPORT:
                pdf_file = io.BytesIO(file_content)
                reader = PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            
            # Word documents
            elif ext in ['.docx'] and PDF_SUPPORT:
                doc_file = io.BytesIO(file_content)
                doc = docx.Document(doc_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text
            
            else:
                return f"[Binary file: {filename} - text extraction not supported for {ext}]"
        
        except Exception as e:
            print(f"⚠ Error extracting text from {filename}: {str(e)}")
            return f"[Error extracting text from {filename}]"
    
    def is_token_expired(self, expiration_time: int) -> bool:
        """Check if the token is expired (with 60 second buffer)"""
        return int(time.time()) >= (expiration_time - 60)
    
    def get_bearer_token(self) -> Optional[str]:
        """Get or refresh the Watson Orchestrate bearer token"""
        if self.wxo_token and not self.is_token_expired(self.wxo_token_expiration):
            return self.wxo_token
        
        try:
            url = "https://iam.platform.saas.ibm.com/siusermgr/api/1.0/apikeys/token"
            response = requests.post(url, json={"apikey": self.wxo_api_key}, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                self.wxo_token = data.get("token")
                self.wxo_token_expiration = int(time.time()) + data.get("expires_in", 600)
                print(f"✓ Watson Orchestrate token obtained (expires in {data.get('expires_in', 600)}s)")
                return self.wxo_token
            else:
                print(f"⚠ Token request failed: {response.status_code} - {response.text}")
                return None
        except Exception as e:
            print(f"⚠ Exception getting token: {str(e)}")
            return None
    
    def trigger_wxo_executive_summary(self, document_text: str, filename: str, doc_id: str, process_id: str) -> Optional[Dict[str, Any]]:
        """Trigger Watson Orchestrate executive summary agent with document content"""
        if not self.wxo_enabled:
            return None
        
        try:
            # Get bearer token
            token = self.get_bearer_token()
            if not token:
                print(f"⚠ Failed to get Watson Orchestrate token")
                return None
            
            # Prepare the prompt for the agent
            process_name = os.getenv('PROCESS_NAME', process_id)
            prompt = f"""Please analyze this document and provide an executive summary:

Document Name: {filename}
Document ID: {doc_id}
Process ID: {process_id}
Process Name: {process_name}

Document Content:
{document_text[:8000]}

Please provide:
1. Executive Summary
2. Key Findings
3. Risk Assessment
4. Recommendations
"""
            
            # Watson Orchestrate chat completions API
            url = f"https://api.dl.watson-orchestrate.ibm.com/instances/{self.wxo_instance_id}/v1/orchestrate/{self.wxo_agent_id}/chat/completions"
            
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "stream": False
            }
            
            print(f"🤖 Triggering Watson Orchestrate executive summary agent...")
            
            response = requests.post(url, headers=headers, json=payload, timeout=120)
            
            if response.status_code == 200:
                data = response.json()
                summary = data.get("choices", [{}])[0].get("message", {}).get("content", "No response returned.")
                print(f"✅ Executive summary generated successfully!")
                print(f"   Summary preview: {summary[:200]}...")
                
                # Use EST timezone for timestamp
                est = pytz.timezone('US/Eastern')
                est_time = datetime.now(est)
                
                result = {
                    "summary": summary,
                    "document_name": filename,
                    "document_id": doc_id,
                    "process_id": process_id,
                    "timestamp": est_time.strftime("%d/%m/%Y %H:%M:%S")
                }
                return result
            else:
                print(f"⚠ Watson Orchestrate agent failed: HTTP {response.status_code}")
                print(f"   Response: {response.text[:500]}")
                return None
        
        except Exception as e:
            print(f"⚠ Exception calling Watson Orchestrate agent: {str(e)}")
            return None
    
    def format_summary_as_docx(self, wxo_result: Dict[str, Any], filename: str, doc_id: str) -> Optional[bytes]:
        """Format Watson Orchestrate summary as a DOCX document"""
        if not PDF_SUPPORT:
            print("⚠ python-docx not available, skipping DOCX generation")
            return None
        
        try:
            from docx import Document
            
            doc = Document()
            
            # Title
            title = doc.add_heading('WATSON ORCHESTRATE EXECUTIVE SUMMARY', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Document Information Section
            doc.add_heading('Document Information', level=1)
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            process_name = os.getenv('PROCESS_NAME', wxo_result.get('process_id', 'N/A'))
            info_data = [
                ('Document Name:', filename),
                ('Document ID:', str(doc_id)),
                ('Process ID:', str(wxo_result.get('process_id', 'N/A'))),
                ('Process Name:', process_name),
                ('Generated:', wxo_result.get('timestamp', 'N/A'))
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[1].text = value
                # Bold the labels
                info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()  # Add spacing
            
            # Executive Summary Section
            doc.add_heading('Executive Summary', level=1)
            
            summary = wxo_result.get('summary', 'No summary available')
            
            # Parse and format the summary
            lines = summary.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    current_paragraph = None
                    continue
                
                # Main headers (starts with ** and ends with ** and no other text)
                if line.startswith('**') and line.endswith('**') and line.count('**') == 2:
                    header_text = line.replace('**', '').strip()
                    doc.add_heading(header_text, level=2)
                    current_paragraph = None
                # Numbered items
                elif len(line) > 2 and line[0:2].replace('.', '').replace(')', '').isdigit():
                    p = doc.add_paragraph(style='List Number')
                    self._add_formatted_text(p, line)
                    current_paragraph = None
                # Bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_text(p, line[2:])
                    current_paragraph = None
                # Regular text
                else:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                        self._add_formatted_text(current_paragraph, line)
                    else:
                        current_paragraph.add_run(' ')
                        self._add_formatted_text(current_paragraph, line)
            
            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            return docx_bytes.getvalue()
        
        except Exception as e:
            print(f"⚠ Error formatting DOCX: {str(e)}")
            return None
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text to paragraph with inline bold formatting support"""
        # Split text by ** markers for bold formatting
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                bold_text = part[2:-2]  # Remove ** markers
                run = paragraph.add_run(bold_text)
                run.font.bold = True
            elif part:
                # Regular text
                paragraph.add_run(part)
    
    async def get_document_folder_id(self, process_id: str) -> Optional[str]:
        """
        Get the folder ID where documents associated with this process are stored.
        This checks existing child documents to find the correct folder.
        """
        try:
            import httpx
            base = self.client.base_url.replace('/openpages', '').rstrip('/')
            url = f"{base}/grc/api/contents/{process_id}/associations/children"
            
            async with httpx.AsyncClient(verify=False, follow_redirects=True) as http_client:
                response = await http_client.get(
                    url,
                    auth=(self.client.username, self.client.password),
                    timeout=30.0
                )
                
                if response.status_code == 200:
                    children = response.json()
                    
                    if not children:
                        print(f"      ⚠ No child documents found for process")
                        return None
                    
                    # Try to find any document to get its folder
                    # First try documents with type 22, then fall back to any child
                    print(f"      🔍 Found {len(children)} child items, checking for documents...")
                    
                    for child in children:
                        doc_id = child.get('id')
                        type_id = child.get('typeDefinitionId')
                        name = child.get('name', 'unknown')
                        
                        print(f"         - Child: {name} (ID: {doc_id}, Type: {type_id})")
                        
                        # Try to get this item's details to find its folder
                        if doc_id:
                            doc_url = f"{base}/grc/api/contents/{doc_id}"
                            doc_response = await http_client.get(
                                doc_url,
                                auth=(self.client.username, self.client.password),
                                timeout=30.0
                            )
                            
                            if doc_response.status_code == 200:
                                doc_data = doc_response.json()
                                folder_id = doc_data.get('parentFolderId')
                                
                                if folder_id:
                                    print(f"      ✓ Found document folder ID: {folder_id} from {name}")
                                    return folder_id
                    
                    print(f"      ⚠ Could not determine folder from any child documents")
                else:
                    print(f"      ⚠ Failed to get process children: HTTP {response.status_code}")
                
                return None
                
        except Exception as e:
            print(f"      ⚠ Exception getting document folder: {str(e)}")
            return None
    async def find_existing_summary(self, process_id: str, original_doc_name: str) -> Optional[str]:
        """
        Find existing executive summary document for a given original document name.
        Returns the document ID if found, None otherwise.
        """
        try:
            import httpx
            base = self.client.base_url.replace('/openpages', '').rstrip('/')
            url = f"{base}/grc/api/contents/{process_id}/associations/children"
            
            # Expected summary name format: original_name-Executive-Summary
            expected_summary_name = f"{original_doc_name}-Executive-Summary"
            
            async with httpx.AsyncClient(verify=False, follow_redirects=True) as http_client:
                response = await http_client.get(
                    url,
                    auth=(self.client.username, self.client.password),
                    timeout=30.0
                )
                
                if response.status_code == 200:
                    children = response.json()
                    
                    # Search for existing summary document
                    for child in children:
                        child_name = child.get('name', '')
                        child_id = child.get('id')
                        
                        # Check if this is the summary document we're looking for
                        if child_name == expected_summary_name or child_name == f"{expected_summary_name}.docx":
                            print(f"      ✓ Found existing summary: {child_name} (ID: {child_id})")
                            return child_id
                    
                    print(f"      ℹ️  No existing summary found for: {expected_summary_name}")
                    return None
                else:
                    print(f"      ⚠ Failed to search for existing summary: HTTP {response.status_code}")
                    return None
                    
        except Exception as e:
            print(f"      ⚠ Exception searching for existing summary: {str(e)}")
            return None
    
    async def delete_document(self, doc_id: str) -> bool:
        """Delete a document from OpenPages by ID"""
        try:
            import httpx
            base = self.client.base_url.replace('/openpages', '').rstrip('/')
            url = f"{base}/grc/api/contents/{doc_id}"
            
            async with httpx.AsyncClient(verify=False, follow_redirects=True) as http_client:
                response = await http_client.delete(
                    url,
                    auth=(self.client.username, self.client.password),
                    timeout=30.0
                )
                
                if response.status_code in [200, 204]:
                    print(f"      ✓ Deleted existing summary (ID: {doc_id})")
                    return True
                else:
                    print(f"      ⚠ Failed to delete document: HTTP {response.status_code}")
                    return False
                    
        except Exception as e:
            print(f"      ⚠ Exception deleting document: {str(e)}")
            return False
    
    
    async def upload_document_to_openpages(self, file_content: bytes, filename: str, description: str, process_id: str) -> Optional[str]:
        """
        Upload a document back to OpenPages and associate it with the process.
        Uses find-and-replace logic: if a summary already exists, delete it first.
        """
        try:
            import httpx
            
            # Get the folder ID where documents for this process are stored
            folder_id = await self.get_document_folder_id(process_id)
            
            if not folder_id:
                print(f"      ⚠ Cannot upload: Unable to determine document folder for process {process_id}")
                return None
            
            # Extract original document name from filename (remove -Executive-Summary suffix)
            # Expected format: "doc_1 [3]-Executive-Summary" or "Executive Risk Summary - doc_1 [3].txt"
            original_doc_name = filename
            if '-Executive-Summary' in filename:
                original_doc_name = filename.split('-Executive-Summary')[0].strip()
            elif filename.startswith('Executive Risk Summary - '):
                original_doc_name = filename.replace('Executive Risk Summary - ', '').replace('.docx', '').replace('.txt', '').strip()
            
            # Check if summary already exists and delete it
            existing_summary_id = await self.find_existing_summary(process_id, original_doc_name)
            if existing_summary_id:
                print(f"      🔄 Replacing existing summary...")
                await self.delete_document(existing_summary_id)
            
            # Encode file content to base64 for binary files
            file_content_b64 = base64.b64encode(file_content).decode('utf-8')
            
            base = self.client.base_url.replace('/openpages', '').rstrip('/')
            
            # Create clean filename: original_name-Executive-Summary (without .docx, OpenPages adds it)
            clean_filename = f"{original_doc_name}-Executive-Summary"
            
            # Create document in the same folder as other process documents
            create_url = f"{base}/grc/api/contents"
            
            create_payload = {
                "contentDefinition": {
                    "attribute": {
                        "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    },
                    "children": file_content_b64
                },
                "fileTypeDefinition": {
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                },
                "fields": {
                    "field": []
                },
                "typeDefinitionId": "4",  # SOXDocument type (same as other documents)
                "parentFolderId": folder_id,
                "name": clean_filename,
                "description": description
            }
            
            async with httpx.AsyncClient(verify=False, follow_redirects=True) as http_client:
                print(f"      📤 Creating document in OpenPages: {clean_filename}")
                create_response = await http_client.post(
                    create_url,
                    json=create_payload,
                    auth=(self.client.username, self.client.password),
                    timeout=60.0
                )
                
                if create_response.status_code not in [200, 201]:
                    print(f"      ⚠ Failed to create document: HTTP {create_response.status_code}")
                    print(f"         Response: {create_response.text[:500]}")
                    return None
                
                doc_data = create_response.json()
                new_doc_id = doc_data.get('id')
                print(f"      ✓ Document created with ID: {new_doc_id}")
                print(f"      ✓ Document automatically associated via parentFolderId: {folder_id}")
                
                return new_doc_id
                
        except Exception as e:
            print(f"      ⚠ Exception uploading document to OpenPages: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    
    def start_health_server(self, port: int = 8080):
        """Start a simple HTTP health check server in a background thread"""
        async def health_handler(request):
            return web.Response(text='OK', status=200)
        
        async def run_server():
            app = web.Application()
            app.router.add_get('/health', health_handler)
            app.router.add_get('/', health_handler)
            runner = web.AppRunner(app)
            await runner.setup()
            site = web.TCPSite(runner, '0.0.0.0', port)
            await site.start()
            print(f"✓ Health check server started on port {port}")
            # Keep the server running
            while True:
                await asyncio.sleep(3600)
        
        def run_in_thread():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            loop.run_until_complete(run_server())
        
        thread = threading.Thread(target=run_in_thread, daemon=True)
        thread.start()
    
    async def monitor_process_for_new_documents(self, process_id: str, check_interval: int = 60):
        """
        Continuously monitor a process for new document uploads
        
        Args:
            process_id: The process ID to monitor
            check_interval: Seconds between checks
        """
        known_documents: set = set()
        
        # Start health check server for Code Engine readiness probe
        self.start_health_server(port=8080)
        
        print("\n" + "=" * 70)
        print("DOCUMENT MONITORING STARTED")
        print("=" * 70)
        print(f"Process ID: {process_id}")
        print(f"Check Interval: {check_interval} seconds")
        print(f"COS Upload: {'ENABLED ☁️' if self.cos_client else 'DISABLED'}")
        print(f"Health Check: ENABLED on port 8080")
        print("=" * 70)
        
        # Initialize - get existing documents
        print("\n📥 Initializing - scanning for existing documents...")
        documents = await self.find_process_documents(process_id)
        
        if documents:
            print(f"   Found {len(documents)} existing document(s)")
            for doc in documents:
                doc_id = self._get_field_value(doc, 'Resource ID')
                doc_name = self._get_field_value(doc, 'Name')
                known_documents.add(doc_id)
                print(f"      - {doc_name} (ID: {doc_id})")
            
            # Download existing documents if COS is enabled
            if self.cos_client:
                print(f"\n   Downloading existing documents to COS...")
                await self.download_documents_to_cos(process_id, documents)
        else:
            print(f"   No existing documents found")
        
        print(f"\n✓ Initialization complete")
        print(f"👀 Now monitoring for NEW documents... (Press Ctrl+C to stop)\n")
        
        # Start monitoring loop
        check_count = 0
        try:
            while True:
                check_count += 1
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                # Check for documents
                current_documents = await self.find_process_documents(process_id)
                
                # Find new documents (skip AI-generated summaries)
                new_docs = []
                for doc in current_documents:
                    doc_id = self._get_field_value(doc, 'Resource ID')
                    doc_name = self._get_field_value(doc, 'Name')
                    
                    # Skip if already tracked
                    if doc_id in known_documents:
                        continue
                    
                    # Skip AI-generated summaries (documents uploaded from COS)
                    if doc_name.startswith("Executive Risk Summary"):
                        print(f"[{timestamp}] ⏭ Skipping AI-generated summary: {doc_name}")
                        known_documents.add(doc_id)  # Track it but don't process
                        continue
                    
                    # This is a new user-uploaded document
                    new_docs.append(doc)
                    known_documents.add(doc_id)
                
                if new_docs:
                    print(f"[{timestamp}] 🆕 NEW DOCUMENT(S) DETECTED: {len(new_docs)}")
                    for doc in new_docs:
                        doc_id = self._get_field_value(doc, 'Resource ID')
                        doc_name = self._get_field_value(doc, 'Name')
                        print(f"   - {doc_name} (ID: {doc_id})")
                    
                    # Download new documents if COS is enabled
                    if self.cos_client:
                        print(f"\n   Downloading new documents to COS...")
                        await self.download_documents_to_cos(process_id, new_docs)
                    
                    print()
                else:
                    if check_count % 10 == 0:  # Log every 10 checks
                        print(f"[{timestamp}] ℹ️  No new documents (tracking {len(known_documents)} documents)")
                
                # Wait for next check
                await asyncio.sleep(check_interval)
                
        except KeyboardInterrupt:
            print(f"\n\n⏹ Monitoring stopped by user")
            print(f"   Total documents tracked: {len(known_documents)}")
    
    async def download_documents_to_cos(self, process_id: str, documents: List[Dict[str, Any]]) -> int:
        """
        Download documents from OpenPages and upload to IBM Cloud Object Storage
        
        Args:
            process_id: The process ID
            documents: List of documents to download
            
        Returns:
            Number of documents successfully downloaded
        """
        if not self.cos_client:
            print("❌ IBM Cloud Object Storage not configured")
            return 0
        
        if not documents:
            print("ℹ️  No documents to download")
            return 0
        
        print(f"\n📥 Downloading {len(documents)} document(s) to COS...")
        success_count = 0
        
        try:
            import httpx
            base = self.client.base_url.replace('/openpages', '').rstrip('/')
            
            async with httpx.AsyncClient(verify=False, follow_redirects=True) as http_client:
                for i, doc in enumerate(documents, 1):
                    doc_id = self._get_field_value(doc, 'Resource ID')
                    doc_name = self._get_field_value(doc, 'Name')
                    
                    try:
                        # Download document from OpenPages
                        download_url = f"{base}/grc/api/contents/{doc_id}/document"
                        headers = {"Accept": "application/octet-stream"}
                        
                        print(f"   [{i}/{len(documents)}] Downloading {doc_name}...")
                        
                        response = await http_client.get(
                            download_url,
                            headers=headers,
                            auth=(self.client.username, self.client.password),
                            timeout=60.0
                        )
                        
                        if response.status_code == 200:
                            file_content = response.content
                            file_size = len(file_content)
                            
                            # Upload original document to COS
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            name, ext = os.path.splitext(doc_name)
                            unique_filename = f"{name}_{timestamp}{ext}"
                            cos_key = f"Process_{process_id}/{unique_filename}"
                            
                            self.cos_client.put_object(
                                Bucket=self.cos_bucket,
                                Key=cos_key,
                                Body=file_content
                            )
                            
                            print(f"      ✅ Uploaded to COS: {cos_key} ({file_size:,} bytes)")
                            
                            # Extract text content
                            print(f"      📄 Extracting text content...")
                            document_text = self.extract_text_from_file(file_content, doc_name)
                            text_length = len(document_text)
                            print(f"      ✓ Extracted {text_length:,} characters of text")
                            
                            # Trigger Watson Orchestrate executive summary agent
                            if self.wxo_enabled and text_length > 50:
                                wxo_result = self.trigger_wxo_executive_summary(document_text, doc_name, doc_id, process_id)
                                if wxo_result:
                                    # Save summary as JSON to COS
                                    summary_filename_json = f"{name}_summary_{timestamp}.json"
                                    summary_content_json = json.dumps(wxo_result, indent=2).encode('utf-8')
                                    cos_key_json = f"Process_{process_id}/{summary_filename_json}"
                                    
                                    self.cos_client.put_object(
                                        Bucket=self.cos_bucket,
                                        Key=cos_key_json,
                                        Body=summary_content_json
                                    )
                                    print(f"      ✅ Summary JSON saved to COS: {cos_key_json}")
                                    
                                    # Format and save summary as DOCX
                                    formatted_doc_bytes = self.format_summary_as_docx(wxo_result, doc_name, doc_id)
                                    if formatted_doc_bytes:
                                        summary_filename_docx = f"{name}_summary_{timestamp}.docx"
                                        cos_key_docx = f"Process_{process_id}/{summary_filename_docx}"
                                        
                                        self.cos_client.put_object(
                                            Bucket=self.cos_bucket,
                                            Key=cos_key_docx,
                                            Body=formatted_doc_bytes
                                        )
                                        print(f"      ✅ Summary DOCX saved to COS: {cos_key_docx}")
                                        
                                        # Upload the DOCX summary back to OpenPages
                                        executive_summary_filename = f"Executive Risk Summary - {doc_name}.docx"
                                        # Use the stored process resource ID instead of process name
                                        resource_id = getattr(self, 'process_resource_id', process_id)
                                        uploaded_doc_id = await self.upload_document_to_openpages(
                                            file_content=formatted_doc_bytes,
                                            filename=executive_summary_filename,
                                            description=f"AI-generated executive risk summary for {doc_name}",
                                            process_id=resource_id
                                        )
                                        
                                        if uploaded_doc_id:
                                            print(f"      ✅ Executive summary uploaded to OpenPages: {executive_summary_filename}")
                                        else:
                                            print(f"      ⚠ Could not upload summary to OpenPages (saved to COS only)")
                                        
                                        print(f"      📄 Executive summary available in COS bucket and OpenPages")
                            
                            success_count += 1
                        else:
                            print(f"      ❌ Failed to download: HTTP {response.status_code}")
                    
                    except Exception as e:
                        print(f"      ❌ Error: {str(e)}")
                        import traceback
                        traceback.print_exc()
                        continue
        
        except Exception as e:
            print(f"❌ Error during download: {str(e)}")
        
        print(f"\n✅ Successfully downloaded {success_count}/{len(documents)} document(s) to COS")
        return success_count
    
    def _get_field_value(self, process: Dict[str, Any], field_name: str) -> str:
        """Extract field value from process data"""
        if 'fields' in process:
            for field in process['fields']:
                if field.get('name') == field_name:
                    value = field.get('value')
                    if value is None:
                        return 'N/A'
                    elif isinstance(value, dict):
                        # Handle enum values
                        return value.get('name', str(value))
                    elif isinstance(value, bool):
                        return 'Yes' if value else 'No'
                    else:
                        return str(value)
        return 'N/A'
    
    def display_process(self, process: Dict[str, Any]):
        """Display process information in a readable format"""
        print("\n" + "=" * 70)
        print("PROCESS DETAILS")
        print("=" * 70)
        
        # Extract key fields
        resource_id = self._get_field_value(process, 'Resource ID')
        name = self._get_field_value(process, 'Name')
        description = self._get_field_value(process, 'Description')
        location = self._get_field_value(process, 'Location')
        status = self._get_field_value(process, 'OPSS-Process:Status')
        business_owner = self._get_field_value(process, 'OPSS-Process:Business Owner')
        process_owner = self._get_field_value(process, 'OPSS-Process:Process Owner')
        created_by = self._get_field_value(process, 'Created By')
        creation_date = self._get_field_value(process, 'Creation Date')
        
        print(f"Resource ID:      {resource_id}")
        print(f"Name:             {name}")
        print(f"Description:      {description}")
        print(f"Location:         {location}")
        print(f"Status:           {status}")
        print(f"Business Owner:   {business_owner}")
        print(f"Process Owner:    {process_owner}")
        print(f"Created By:       {created_by}")
        print(f"Creation Date:    {creation_date}")
        
        # Display all fields if available
        if 'fields' in process:
            print("\nAll Fields:")
            for field in process['fields']:
                field_name = field.get('name', 'Unknown')
                value = field.get('value')
                if value is not None:
                    if isinstance(value, dict):
                        value_str = value.get('name', str(value))
                    else:
                        value_str = str(value)
                    # Truncate long values
                    if len(value_str) > 60:
                        value_str = value_str[:60] + "..."
                    print(f"  {field_name}: {value_str}")
        
        print("=" * 70 + "\n")
    
    def display_process_list(self, processes: List[Dict[str, Any]]):
        """Display a list of processes in a table format"""
        if not processes:
            print("\nNo processes to display.\n")
            return
        
        print("\n" + "=" * 120)
        print(f"{'Resource ID':<15} {'Name':<25} {'Description':<40} {'Status':<20}")
        print("=" * 120)
        
        for process in processes:
            resource_id = self._get_field_value(process, 'Resource ID')
            name = self._get_field_value(process, 'Name')
            description = self._get_field_value(process, 'Description')
            status = self._get_field_value(process, 'OPSS-Process:Status')
            
            # Truncate long values
            if len(name) > 22:
                name = name[:22] + "..."
            if len(description) > 37:
                description = description[:37] + "..."
            if len(status) > 17:
                status = status[:17] + "..."
            
            print(f"{resource_id:<15} {name:<25} {description:<40} {status:<20}")
        
        print("=" * 120 + "\n")
    
    def display_documents(self, documents: List[Dict[str, Any]]):
        """Display a list of documents in a table format"""
        if not documents:
            print("\nNo documents to display.\n")
            return
        
        print("\n" + "=" * 120)
        print(f"{'Resource ID':<15} {'Name':<40} {'Created':<25} {'Location':<40}")
        print("=" * 120)
        
        for doc in documents:
            resource_id = self._get_field_value(doc, 'Resource ID')
            name = self._get_field_value(doc, 'Name')
            created = self._get_field_value(doc, 'Creation Date')
            location = self._get_field_value(doc, 'Location')
            
            # Truncate long values
            if len(name) > 37:
                name = name[:37] + "..."
            if len(created) > 22:
                created = created[:22] + "..."
            if len(location) > 37:
                location = location[:37] + "..."
            
            print(f"{resource_id:<15} {name:<40} {created:<25} {location:<40}")
        
        print("=" * 120 + "\n")


async def main():
    """Main function"""
    parser = argparse.ArgumentParser(
        description='Find OpenPages processes by ID or pattern',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Find a specific process by ID
  python find_process.py --id AML_PROC_0008
  
  # Find processes matching a pattern
  python find_process.py --pattern "AML_PROC_*"
  
  # List all processes (up to 50)
  python find_process.py --all
  
  # List all processes with custom limit
  python find_process.py --all --limit 100
        """
    )
    
    parser.add_argument('--id', type=str, help='Find process by exact ID (e.g., AML_PROC_0008)')
    parser.add_argument('--pattern', type=str, help='Find processes matching pattern (e.g., AML_PROC_*)')
    parser.add_argument('--all', action='store_true', help='List all processes')
    parser.add_argument('--documents', action='store_true', help='Show documents (uploaded files) for the process')
    parser.add_argument('--download', action='store_true', help='Download documents to IBM Cloud Object Storage')
    parser.add_argument('--monitor', action='store_true', help='Continuously monitor for new document uploads')
    parser.add_argument('--interval', type=int, default=60, help='Monitor check interval in seconds (default: 60)')
    parser.add_argument('--limit', type=int, default=50, help='Maximum number of processes to return (default: 50)')
    parser.add_argument('--verbose', '-v', action='store_true', help='Show detailed process information')
    
    args = parser.parse_args()
    
    # Validate arguments
    if not any([args.id, args.pattern, args.all]):
        parser.error('Must specify one of: --id, --pattern, or --all')
    
    # Initialize OpenPages client
    print("🔌 Connecting to OpenPages...")
    print(f"   Server: {settings.OPENPAGES_BASE_URL}")
    print(f"   Auth Type: {settings.OPENPAGES_AUTHENTICATION_TYPE}")
    
    try:
        client = OpenPagesClient(
            base_url=settings.OPENPAGES_BASE_URL,
            auth_type=settings.OPENPAGES_AUTHENTICATION_TYPE,
            username=settings.OPENPAGES_USERNAME,
            password=settings.OPENPAGES_PASSWORD,
            api_key=settings.OPENPAGES_APIKEY,
            authentication_url=settings.OPENPAGES_AUTHENTICATION_URL
        )
        
        # Initialize authentication
        await client.initialize_auth()
        print("✅ Connected successfully\n")
        
        # Create finder
        finder = ProcessFinder(client)
        
        # Execute search based on arguments
        if args.id:
            # If --monitor flag is set, start continuous monitoring
            if args.monitor:
                await finder.monitor_process_for_new_documents(args.id, args.interval)
            else:
                # Find by exact ID
                process = await finder.find_process_by_id(args.id)
                if process:
                    if args.verbose:
                        finder.display_process(process)
                    else:
                        finder.display_process_list([process])
                    
                    # If --documents flag is set, also show documents
                    if args.documents:
                        print("\n" + "=" * 70)
                        print("DOCUMENTS IN PROCESS")
                        print("=" * 70)
                        documents = await finder.find_process_documents(args.id)
                        if documents:
                            finder.display_documents(documents)
                            
                            # If --download flag is set, download documents to COS
                            if args.download:
                                await finder.download_documents_to_cos(args.id, documents)
        
        elif args.pattern:
            # Find by pattern
            processes = await finder.find_processes_by_pattern(args.pattern)
            if processes:
                if args.verbose:
                    for process in processes:
                        finder.display_process(process)
                else:
                    finder.display_process_list(processes)
        
        elif args.all:
            # Find all processes
            processes = await finder.find_all_processes(limit=args.limit)
            if processes:
                if args.verbose:
                    for process in processes:
                        finder.display_process(process)
                else:
                    finder.display_process_list(processes)
        
        # Close client
        await client.close()
        print("✅ Done!")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())

# Made with Bob
